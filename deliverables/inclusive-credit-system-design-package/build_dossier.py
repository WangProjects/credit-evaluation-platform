from __future__ import annotations

import json
import re
import zipfile
from pathlib import Path
from tempfile import NamedTemporaryFile
from xml.etree import ElementTree

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
OUT = ROOT / "Inclusive_Credit_Platform_Technical_Dossier.docx"
SOURCES = [
    ROOT / "README.md",
    ROOT / "01_SYSTEM_ARCHITECTURE.md",
    ROOT / "02_TECHNICAL_SPECIFICATIONS.md",
    ROOT / "03_COMPONENT_DOCUMENTATION.md",
    ROOT / "04_MODEL_DEVELOPMENT_AND_TESTING.md",
    ROOT / "05_FAIRNESS_AND_BIAS_TESTING.md",
    ROOT / "06_EVIDENCE_INDEX.md",
    ROOT / "07_RISK_REGISTER_AND_ROADMAP.md",
    ROOT / "08_ARCHITECTURE_DIAGRAMS.md",
    ROOT / "09_PRODUCT_DEMONSTRATION.md",
    ROOT / "10_SUBMISSION_READINESS.md",
]
EVIDENCE_ROOT = ROOT / "evidence"
DIAGRAM_ROOT = ROOT / "diagrams"
DIAGRAM_DEFINITIONS = [
    ("System context", DIAGRAM_ROOT / "system_context.mmd"),
    ("Online scoring sequence", DIAGRAM_ROOT / "scoring_sequence.mmd"),
    ("Production deployment topology", DIAGRAM_ROOT / "production_deployment_topology.mmd"),
    ("Model lifecycle and governance", DIAGRAM_ROOT / "model_lifecycle_governance.mmd"),
    ("Data lineage and evidence", DIAGRAM_ROOT / "data_lineage_evidence.mmd"),
    ("Fairness control loop", DIAGRAM_ROOT / "fairness_control_loop.mmd"),
]
CORE_NS = {
    "dc": "http://purl.org/dc/elements/1.1/",
    "dcterms": "http://purl.org/dc/terms/",
}

NAVY = RGBColor(27, 54, 93)
BLUE = RGBColor(46, 116, 181)
SLATE = RGBColor(80, 90, 105)
LIGHT = "E8EEF5"


def font(run, size=11, bold=False, italic=False, color=None):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_table_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    cant_split.set(qn("w:val"), "true")
    tr_pr.append(cant_split)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    font(run, 9, color=SLATE)
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char)
    run._r.append(instr)
    run._r.append(fld_end)


def configure(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.82)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.88)
    section.right_margin = Inches(0.88)
    section.header_distance = Inches(0.36)
    section.footer_distance = Inches(0.36)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.08

    for name, size, before, after, color in (
        ("Heading 1", 17, 16, 8, NAVY),
        ("Heading 2", 13.5, 12, 6, BLUE),
        ("Heading 3", 11.5, 8, 4, NAVY),
    ):
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for sec in doc.sections:
        hp = sec.header.paragraphs[0]
        hp.text = "INCLUSIVE CREDIT PLATFORM  |  TECHNICAL EVIDENCE DOSSIER"
        hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        font(hp.runs[0], 8.5, bold=True, color=SLATE)
        add_page_number(sec.footer.paragraphs[0])


def add_cover(doc):
    for _ in range(5):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    font(p.add_run("TECHNICAL EVIDENCE DOSSIER"), 12, bold=True, color=BLUE)
    p.paragraph_format.space_after = Pt(18)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    font(p.add_run("Inclusive Credit Evaluation Platform"), 28, bold=True, color=NAVY)
    p.paragraph_format.space_after = Pt(10)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    font(p.add_run("Architecture, specifications, model development, fairness testing, and implementation evidence"), 14, color=SLATE)
    p.paragraph_format.space_after = Pt(60)
    for label, value in (
        ("Prepared by", "Mr. Wang"),
        ("Repository", "credit-evaluation-platform"),
        ("Document purpose", "Technical review and implementation evidence"),
        ("Evidence status", "Reference implementation; synthetic-data demonstration; not production approved"),
    ):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        font(p.add_run(f"{label}: "), 10, bold=True, color=NAVY)
        font(p.add_run(value), 10, color=SLATE)
    doc.add_page_break()


def add_contents(doc):
    p = doc.add_heading("Contents", level=1)
    p.paragraph_format.space_before = Pt(0)
    items = [
        "Executive package overview",
        "1. System Architecture",
        "2. Technical Specifications",
        "3. Component Documentation",
        "4. Model Development and Testing",
        "5. Fairness and Bias Testing",
        "6. Evidence Index",
        "7. Risk Register and Completion Roadmap",
        "8. Architecture Diagram Catalog",
        "9. Product Demonstration and Operator Walkthrough",
        "10. Submission Readiness, Traceability, and Acceptance Package",
        "Appendix A. Test and Training Evidence",
        "Appendix B. Baseline Training Report",
        "Appendix C. Model Registry Snapshot",
        "Appendix D. Editable Diagram Definitions",
        "Appendix E. Screenshot Integrity Manifest",
    ]
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(5)
        font(p.add_run(item), 11)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    font(p.add_run("Evidence convention. "), 10.5, bold=True, color=NAVY)
    font(p.add_run("Implemented work, partial scaffolding, and planned controls are labeled separately throughout the dossier."), 10.5)
    doc.add_page_break()


def add_inline(paragraph, text):
    parts = re.split(r"(`[^`]+`|\*\*[^*]+\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            font(run, 9.2, color=NAVY)
            run.font.name = "Courier New"
        elif part.startswith("**") and part.endswith("**"):
            font(paragraph.add_run(part[2:-2]), 10.5, bold=True)
        else:
            font(paragraph.add_run(part), 10.5)


def add_table(doc, rows):
    parsed = [[c.strip() for c in row.strip().strip("|").split("|")] for row in rows]
    if len(parsed) > 1 and all(set(c) <= set("-: ") for c in parsed[1]):
        parsed.pop(1)
    cols = max(len(r) for r in parsed)
    table = doc.add_table(rows=0, cols=cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = [Inches(6.74 / cols)] * cols
    for ridx, row in enumerate(parsed):
        cells = table.add_row().cells
        for idx in range(cols):
            cell = cells[idx]
            cell.width = widths[idx]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            set_cell_margins(cell)
            text = row[idx] if idx < len(row) else ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            add_inline(p, text)
            for run in p.runs:
                run.font.size = Pt(8.6)
                if ridx == 0:
                    run.bold = True
                    run.font.color.rgb = NAVY
            if ridx == 0:
                shade(cell, LIGHT)
        if ridx == 0:
            set_repeat_table_header(table.rows[0])
        prevent_table_row_split(table.rows[-1])
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_code(doc, lines):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    shade(cell, "F5F7FA")
    set_cell_margins(cell, top=100, bottom=100, start=140, end=140)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("\n".join(lines))
    run.font.name = "Courier New"
    run.font.size = Pt(8.2)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_figure(doc, image_path, caption):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_together = True
    run = p.add_run()
    picture = run.add_picture(str(image_path), width=Inches(6.65))
    picture._inline.docPr.set("descr", caption)
    picture._inline.docPr.set("title", caption)
    cp = doc.add_paragraph()
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp.paragraph_format.space_before = Pt(2)
    cp.paragraph_format.space_after = Pt(8)
    cp.paragraph_format.keep_with_next = True
    font(cp.add_run(caption), 9, italic=True, color=SLATE)


def new_numbering_id(doc):
    numbering = doc.part.numbering_part.element
    existing = [int(n.get(qn("w:numId"))) for n in numbering.findall(qn("w:num"))]
    num_id = max(existing, default=0) + 1
    abstract_id = 0
    style = doc.styles["List Number"]
    num_pr = style.element.pPr.numPr if style.element.pPr is not None else None
    if num_pr is not None and num_pr.numId is not None:
        base_num_id = num_pr.numId.val
        base_num = next((n for n in numbering.findall(qn("w:num")) if int(n.get(qn("w:numId"))) == base_num_id), None)
        if base_num is not None:
            abstract_id = int(base_num.find(qn("w:abstractNumId")).get(qn("w:val")))
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract = OxmlElement("w:abstractNumId")
    abstract.set(qn("w:val"), str(abstract_id))
    num.append(abstract)
    override = OxmlElement("w:lvlOverride")
    override.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:startOverride")
    start.set(qn("w:val"), "1")
    override.append(start)
    num.append(override)
    numbering.append(num)
    return num_id


def apply_num_id(paragraph, num_id):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.get_or_add_numPr()
    num_pr.get_or_add_ilvl().val = 0
    num_pr.get_or_add_numId().val = num_id


def append_markdown(doc, path, first=False):
    lines = path.read_text(encoding="utf-8").splitlines()
    if not first:
        doc.add_page_break()
    i = 0
    diagram_view_seen = False
    demonstration_view_seen = False
    while i < len(lines):
        line = lines[i].rstrip()
        if line.startswith("```"):
            block = []
            i += 1
            while i < len(lines) and not lines[i].startswith("```"):
                block.append(lines[i])
                i += 1
            add_code(doc, block)
        elif line.startswith("|"):
            rows = []
            while i < len(lines) and lines[i].startswith("|"):
                rows.append(lines[i])
                i += 1
            add_table(doc, rows)
            continue
        elif line.startswith("!["):
            match = re.match(r"!\[([^]]+)\]\(([^)]+)\)", line)
            if match:
                caption, relative_path = match.groups()
                add_figure(doc, path.parent / relative_path, caption)
        elif line.startswith("# "):
            p = doc.add_heading(line[2:].strip(), level=1)
            p.paragraph_format.space_before = Pt(0)
            for run in p.runs:
                font(run, 17, bold=True, color=NAVY)
        elif line.startswith("## "):
            if line.startswith("## Architecture view"):
                if diagram_view_seen:
                    doc.add_page_break()
                diagram_view_seen = True
            if line.startswith("## Demonstration screen"):
                if demonstration_view_seen:
                    doc.add_page_break()
                demonstration_view_seen = True
            p = doc.add_heading(line[3:].strip(), level=2)
            for run in p.runs:
                font(run, 13.5, bold=True, color=BLUE)
        elif line.startswith("### "):
            p = doc.add_heading(line[4:].strip(), level=3)
            for run in p.runs:
                font(run, 11.5, bold=True, color=NAVY)
        elif re.match(r"^\d+\. ", line):
            num_id = new_numbering_id(doc)
            while i < len(lines) and re.match(r"^\d+\. ", lines[i].rstrip()):
                p = doc.add_paragraph(style="List Number")
                apply_num_id(p, num_id)
                add_inline(p, re.sub(r"^\d+\. ", "", lines[i].rstrip()))
                i += 1
            continue
        elif line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_inline(p, line[2:])
        elif line.strip():
            p = doc.add_paragraph()
            add_inline(p, line)
        i += 1


def add_json_appendix(doc, title, path):
    doc.add_page_break()
    heading = doc.add_heading(title, level=1)
    heading.paragraph_format.space_before = Pt(0)
    payload = json.loads(path.read_text(encoding="utf-8"))
    p = doc.add_paragraph()
    add_inline(
        p,
        "Machine-readable evidence is reproduced below without interpretation so the printed record can be reconciled with the repository artifact.",
    )
    add_code(doc, json.dumps(payload, indent=2, sort_keys=True).splitlines())


def add_text_appendix(doc, title, path, introduction):
    doc.add_page_break()
    heading = doc.add_heading(title, level=1)
    heading.paragraph_format.space_before = Pt(0)
    p = doc.add_paragraph()
    add_inline(p, introduction)
    add_code(doc, path.read_text(encoding="utf-8").splitlines())


def add_evidence_appendices(doc):
    append_markdown(doc, EVIDENCE_ROOT / "TEST_RESULTS.md")
    first_heading = next(
        (paragraph for paragraph in reversed(doc.paragraphs) if paragraph.style.name == "Heading 1"),
        None,
    )
    if first_heading is not None:
        first_heading.text = "Appendix A. Test and Training Evidence"
        for run in first_heading.runs:
            font(run, 17, bold=True, color=NAVY)

    add_json_appendix(
        doc,
        "Appendix B. Baseline Training Report",
        EVIDENCE_ROOT / "latest_train_report.json",
    )
    add_json_appendix(
        doc,
        "Appendix C. Model Registry Snapshot",
        EVIDENCE_ROOT / "model_registry.json",
    )

    doc.add_page_break()
    heading = doc.add_heading("Appendix D. Editable Diagram Definitions", level=1)
    heading.paragraph_format.space_before = Pt(0)
    p = doc.add_paragraph()
    add_inline(
        p,
        "The rendered diagrams in Section 8 are the print views. The Mermaid definitions below are the complete editable graph sources used to regenerate those views.",
    )
    for index, (label, path) in enumerate(DIAGRAM_DEFINITIONS):
        if index:
            doc.add_page_break()
        heading = doc.add_heading(label, level=2)
        heading.paragraph_format.space_before = Pt(0)
        p = doc.add_paragraph()
        font(p.add_run(f"Source file: diagrams/{path.name}"), 9.5, italic=True, color=SLATE)
        add_code(doc, path.read_text(encoding="utf-8").splitlines())

    add_text_appendix(
        doc,
        "Appendix E. Screenshot Integrity Manifest",
        EVIDENCE_ROOT / "SCREENSHOT_MANIFEST.sha256",
        "The SHA-256 values below identify the twelve local mock-mode product captures embedded in Section 9. Verify them from the package root with `shasum -a 256 -c evidence/SCREENSHOT_MANIFEST.sha256`.",
    )


def scrub_core_timestamps(path):
    with zipfile.ZipFile(path, "r") as source:
        core = ElementTree.fromstring(source.read("docProps/core.xml"))
        for qualified_name in (
            f"{{{CORE_NS['dcterms']}}}created",
            f"{{{CORE_NS['dcterms']}}}modified",
            f"{{{CORE_NS['dc']}}}description",
        ):
            node = core.find(qualified_name)
            if node is not None:
                core.remove(node)
        core_xml = ElementTree.tostring(core, encoding="utf-8", xml_declaration=True)

        with NamedTemporaryFile(dir=path.parent, suffix=".docx", delete=False) as temporary:
            temporary_path = Path(temporary.name)
        try:
            with zipfile.ZipFile(temporary_path, "w") as target:
                for item in source.infolist():
                    payload = core_xml if item.filename == "docProps/core.xml" else source.read(item.filename)
                    target.writestr(item, payload)
            temporary_path.replace(path)
        finally:
            if temporary_path.exists():
                temporary_path.unlink()


def main():
    doc = Document()
    configure(doc)
    add_cover(doc)
    add_contents(doc)
    for index, source in enumerate(SOURCES):
        append_markdown(doc, source, first=index == 0)
    add_evidence_appendices(doc)
    props = doc.core_properties
    props.title = "Inclusive Credit Evaluation Platform — Technical Evidence Dossier"
    props.subject = "Architecture, specifications, model development, fairness testing, and implementation evidence"
    props.author = "Mr. Wang"
    props.keywords = "credit, alternative data, responsible AI, fairness, model risk, architecture"
    doc.save(OUT)
    scrub_core_timestamps(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
